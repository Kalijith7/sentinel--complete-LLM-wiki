"""
wiki_engine/parser.py
=====================
Document parser — Phase 1 of the pipeline.

Parses raw documents into structured dicts that the LLM compiler can work with.
Preserves tables, section headings, and embedded assets.

Intel extension: _parse_docx() extracts embedded images and associates each
image with the caption text and surrounding paragraphs, enabling personnel
identification without requiring a vision model.

Supported formats:
  .pdf            Docling (structure-aware) → PyMuPDF → pypdf
  .docx / .doc    Docling → python-docx (with image extraction)
  .md / .markdown Native Markdown parser
  .txt            Plain text
  .csv            CSV with table structure
  .jpg/.png/...   Tesseract OCR
"""
from __future__ import annotations

import re
import io
import logging
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)


class DocumentParser:
    """
    Parse raw documents into structured dicts for the LLM compiler.

    All parsers return a consistent dict:
        text          str   — full text content (Markdown for structured docs)
        tables        list  — extracted tables as {headers, rows}
        sections      list  — {heading, level} pairs
        filename      str   — original filename
        filetype      str   — detected type
        image_assets  list  — extracted image assets (docx only, others [])

    image_assets entries:
        img_ref       str   — sequential reference "img-0", "img-1" …
        filename      str   — original name in docx zip
        ext           str   — file extension (.png, .jpg, .jpeg …)
        bytes         bytes — raw image data
        caption       str   — verbatim caption / adjacent text
        context       str   — 3 paragraphs before + 3 after
        para_index    int   — paragraph index in document
    """

    async def parse(self, path: Path) -> dict[str, Any]:
        """Route to the correct parser based on file suffix."""
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return await self._parse_pdf(path)
        elif suffix in (".docx", ".doc"):
            return await self._parse_docx(path)
        elif suffix in (".md", ".markdown"):
            return self._parse_markdown(path)
        elif suffix == ".txt":
            return self._parse_text(path)
        elif suffix == ".csv":
            return self._parse_csv(path)
        elif suffix in (".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".gif"):
            return await self._parse_image(path)
        else:
            return self._parse_text(path)

    # ─── PDF ──────────────────────────────────────────────────────────────

    async def _parse_pdf(self, path: Path) -> dict[str, Any]:
        try:
            from docling.document_converter import DocumentConverter
            from docling.datamodel.pipeline_options import PdfPipelineOptions

            opts = PdfPipelineOptions()
            opts.do_ocr = True
            opts.do_table_structure = True

            converter = DocumentConverter()
            result    = converter.convert(str(path))
            doc       = result.document

            return {
                "text":         doc.export_to_markdown(),
                "tables":       self._extract_docling_tables(doc),
                "sections":     self._extract_docling_sections(doc),
                "filename":     path.name,
                "filetype":     "pdf",
                "image_assets": [],
            }
        except ImportError:
            return self._pymupdf_fallback(path)
        except Exception as e:
            log.warning(f"Docling PDF failed for {path.name}: {e}")
            return self._pymupdf_fallback(path)

    def _pymupdf_fallback(self, path: Path) -> dict[str, Any]:
        try:
            import fitz  # PyMuPDF
            doc   = fitz.open(str(path))
            pages = [page.get_text() for page in doc]
            text  = "\n\n".join(pages)
            return {"text": text, "tables": [], "sections": [],
                    "filename": path.name, "filetype": "pdf", "image_assets": []}
        except ImportError:
            pass
        except Exception as e:
            log.warning(f"PyMuPDF failed for {path.name}: {e}")

        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            text   = "\n".join(p.extract_text() or "" for p in reader.pages)
            return {"text": text, "tables": [], "sections": [],
                    "filename": path.name, "filetype": "pdf", "image_assets": []}
        except Exception as e:
            log.warning(f"pypdf failed for {path.name}: {e}")
            return self._empty(path, "pdf")

    # ─── DOCX ─────────────────────────────────────────────────────────────

    async def _parse_docx(self, path: Path) -> dict[str, Any]:
        """
        Parse a .docx file.  Extracts text, tables, and embedded images.

        Image extraction uses python-docx XML relationships to locate
        embedded images, then harvests the caption/adjacent text for each.
        Falls back gracefully — image_assets returns [] on any failure.
        """
        text:   str              = ""
        tables: list[dict]       = []
        images: list[dict]       = []

        # Try Docling first for rich structure
        try:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result    = converter.convert(str(path))
            doc       = result.document
            text      = doc.export_to_markdown()
            tables    = self._extract_docling_tables(doc)
        except Exception as e:
            log.debug(f"Docling docx failed ({path.name}): {e} — using python-docx")
            # Fallback: python-docx for text extraction
            try:
                from docx import Document as DocxDocument
                docx_doc = DocxDocument(str(path))
                paragraphs = [p.text for p in docx_doc.paragraphs]
                text       = "\n".join(paragraphs)
                for t in docx_doc.tables:
                    rows = [[c.text for c in row.cells] for row in t.rows]
                    tables.append({"headers": rows[0] if rows else [], "rows": rows[1:]})
            except Exception as e2:
                log.warning(f"python-docx failed ({path.name}): {e2}")
                text = ""

        # Always attempt image extraction separately via python-docx XML
        images = self._extract_docx_images(path)

        return {
            "text":         text,
            "tables":       tables,
            "sections":     [],
            "filename":     path.name,
            "filetype":     "docx",
            "image_assets": images,
        }

    def _extract_docx_images(self, path: Path) -> list[dict[str, Any]]:
        """
        Extract embedded images from a .docx file using python-docx XML.

        Strategy:
          1. Walk document paragraphs with their index.
          2. For each paragraph, inspect the raw XML for <a:blip> elements
             (inline image anchors).
          3. Resolve the r:embed relationship ID to the image part.
          4. Read raw image bytes.
          5. Determine caption: paragraph's own text + next paragraph
             (if ≤200 chars — likely a figure caption line).
          6. Determine context: 3 paragraphs before + 3 after.
          7. Return list of image dicts with img_ref, bytes, caption, context.

        Returns [] on any failure — never raises.
        """
        results: list[dict[str, Any]] = []

        try:
            from docx import Document as DocxDocument
            from lxml import etree  # type: ignore

            docx_doc   = DocxDocument(str(path))
            paragraphs = docx_doc.paragraphs
            n_paras    = len(paragraphs)

            img_idx = 0

            # XML namespaces used in docx inline images
            NS_A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
            NS_R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            NS_WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            NS_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"

            for para_idx, para in enumerate(paragraphs):
                # Search for <a:blip> in the paragraph XML — indicates an inline image
                blip_elements = para._p.findall(
                    f".//{{{NS_A}}}blip"
                )
                if not blip_elements:
                    # Also check alternate namespace for compatibility
                    blip_elements = para._p.findall(
                        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                    )

                for blip in blip_elements:
                    # Get the embed relationship ID
                    embed_id = blip.get(f"{{{NS_R}}}embed") or blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if not embed_id:
                        continue

                    # Resolve relationship to image part
                    try:
                        rel   = docx_doc.part.rels[embed_id]
                        img_bytes: bytes = rel.target_part.blob

                        # Determine extension from content type
                        ct  = rel.target_part.content_type or ""
                        ext = self._mime_to_ext(ct)
                        orig_name = Path(rel.target_part.partname).name

                    except (KeyError, AttributeError) as rel_err:
                        log.debug(f"Could not resolve image rel {embed_id}: {rel_err}")
                        continue

                    # Determine caption: own paragraph text + next paragraph
                    own_text  = para.text.strip()
                    next_text = ""
                    if para_idx + 1 < n_paras:
                        candidate = paragraphs[para_idx + 1].text.strip()
                        if len(candidate) <= 300:
                            next_text = candidate

                    caption = " | ".join(
                        t for t in [own_text, next_text] if t
                    )

                    # Context window: 3 before + 3 after
                    ctx_start = max(0, para_idx - 3)
                    ctx_end   = min(n_paras, para_idx + 4)
                    context   = " ".join(
                        paragraphs[i].text.strip()
                        for i in range(ctx_start, ctx_end)
                        if paragraphs[i].text.strip()
                    )

                    results.append({
                        "img_ref":    f"img-{img_idx}",
                        "filename":   orig_name,
                        "ext":        ext,
                        "bytes":      img_bytes,
                        "caption":    caption,
                        "context":    context[:600],
                        "para_index": para_idx,
                    })
                    img_idx += 1

            log.info(f"  Extracted {img_idx} image(s) from {path.name}")

        except ImportError:
            log.debug("lxml or python-docx not available — skipping image extraction")
        except Exception as e:
            log.warning(f"Image extraction failed for {path.name}: {e}")

        return results

    @staticmethod
    def _mime_to_ext(content_type: str) -> str:
        """Map MIME content-type to file extension."""
        mapping = {
            "image/png":  ".png",
            "image/jpeg": ".jpg",
            "image/jpg":  ".jpg",
            "image/gif":  ".gif",
            "image/bmp":  ".bmp",
            "image/tiff": ".tiff",
            "image/webp": ".webp",
            "image/emf":  ".emf",
            "image/wmf":  ".wmf",
        }
        return mapping.get(content_type.lower().split(";")[0].strip(), ".png")

    # ─── Markdown ─────────────────────────────────────────────────────────

    def _parse_markdown(self, path: Path) -> dict[str, Any]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        sections = [
            {"heading": m.group(2), "level": len(m.group(1))}
            for m in re.finditer(r"^(#{1,6})\s+(.+)$", text, re.MULTILINE)
        ]
        return {"text": text, "tables": [], "sections": sections,
                "filename": path.name, "filetype": "markdown", "image_assets": []}

    # ─── Plain text ────────────────────────────────────────────────────────

    def _parse_text(self, path: Path) -> dict[str, Any]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return {"text": text, "tables": [], "sections": [],
                "filename": path.name, "filetype": "text", "image_assets": []}

    # ─── CSV ──────────────────────────────────────────────────────────────

    def _parse_csv(self, path: Path) -> dict[str, Any]:
        import csv
        rows: list[list[str]] = []
        with open(path, newline="", encoding="utf-8", errors="ignore") as f:
            for row in csv.reader(f):
                rows.append(row)
                if len(rows) >= 201:
                    break
        headers = rows[0] if rows else []
        text    = "\n".join(",".join(r) for r in rows[:200])
        table   = {"headers": headers, "rows": rows[1:100]}
        return {"text": text, "tables": [table], "sections": [],
                "filename": path.name, "filetype": "csv", "image_assets": []}

    # ─── Image (OCR) ──────────────────────────────────────────────────────

    async def _parse_image(self, path: Path) -> dict[str, Any]:
        """
        Parse a standalone image file using Tesseract OCR.
        Returns text content and marks as image filetype.
        No vision model — caption/context for personnel comes from
        docx image extraction, not standalone image files.
        """
        try:
            import pytesseract
            from PIL import Image
            img  = Image.open(str(path))
            text = pytesseract.image_to_string(img)
            return {"text": text, "tables": [], "sections": [],
                    "filename": path.name, "filetype": "image", "image_assets": []}
        except ImportError:
            return {"text": f"[Image: {path.name} — OCR unavailable]",
                    "tables": [], "sections": [],
                    "filename": path.name, "filetype": "image", "image_assets": []}
        except Exception as e:
            log.warning(f"OCR failed for {path.name}: {e}")
            return {"text": f"[Image: {path.name} — OCR failed]",
                    "tables": [], "sections": [],
                    "filename": path.name, "filetype": "image", "image_assets": []}

    # ─── Docling helpers ──────────────────────────────────────────────────

    @staticmethod
    def _extract_docling_tables(doc: Any) -> list[dict[str, Any]]:
        tables: list[dict[str, Any]] = []
        for table in getattr(doc, "tables", []):
            try:
                rows = [[cell.text for cell in row] for row in table.data.grid]
                tables.append({
                    "headers": rows[0] if rows else [],
                    "rows":    rows[1:] if len(rows) > 1 else [],
                })
            except Exception:
                pass
        return tables

    @staticmethod
    def _extract_docling_sections(doc: Any) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []
        for item in getattr(doc, "texts", []):
            label = str(getattr(item, "label", "")).lower()
            if any(k in label for k in ("section", "title", "heading")):
                sections.append({
                    "heading": getattr(item, "text", ""),
                    "level":   getattr(item, "level", 1),
                })
        return sections

    @staticmethod
    def _empty(path: Path, filetype: str) -> dict[str, Any]:
        return {"text": "", "tables": [], "sections": [],
                "filename": path.name, "filetype": filetype, "image_assets": []}
